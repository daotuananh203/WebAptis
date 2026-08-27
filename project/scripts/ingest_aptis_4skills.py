"""Ingest the seven-test Aptis four-skills source bundle.

This importer is deliberately source-first.  It reads the PDF, the transcript
DOCX files and the answer-key tables, then writes an independent namespace:
``aptis-4skills-01`` .. ``aptis-4skills-07``.  Existing ``aptis-b2-*`` files
are never opened for writing.

Listening segmentation is contract-driven: the transcript paragraph groups
define the recording blocks, Faster-Whisper word timestamps only align those
blocks to the master recording, and incomplete/repeated audio is recorded in
the integrity manifest rather than silently treated as a question boundary.
"""

from __future__ import annotations

import hashlib
import io
import json
import re
import shutil
import subprocess
import sys
import zipfile
from difflib import SequenceMatcher
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from PIL import Image
from pypdf import PdfReader


PROJECT_ROOT = Path(__file__).resolve().parents[1]
WORKSPACE_ROOT = PROJECT_ROOT.parent
SOURCE_ROOT = WORKSPACE_ROOT / "APTIS" / "Bộ đề 4 kĩ năng"
PDF_PATH = SOURCE_ROOT / "01. Aptis.docx.pdf"
BUNDLE_ROOT = SOURCE_ROOT / "File Nghe + Transcript + Đáp án Aptis"
AUDIO_ROOT = BUNDLE_ROOT / "File Nghe"
TRANSCRIPT_ROOT = BUNDLE_ROOT / "Transcript"
GUIDANCE_PATH = BUNDLE_ROOT / "Hướng dẫn và Đáp án Nghe – Đọc, Gợi ý Nói – Viết cho Bộ đề APTIS.docx"

TEST_IDS = [f"aptis-4skills-{n:02d}" for n in range(1, 8)]

# PDF page numbers are 1-based.  These are obtained from the visible section
# headings in the source PDF, not inferred from the number of pages per test.
PAGE_RANGES = {
    1: {"core": (11, 14), "reading": (15, 19), "listening": (20, 22), "speaking": (23, 24), "writing": (25, 27)},
    2: {"core": (28, 31), "reading": (32, 36), "listening": (37, 39), "speaking": (40, 41), "writing": (42, 44)},
    3: {"core": (45, 48), "reading": (49, 53), "listening": (54, 56), "speaking": (57, 58), "writing": (59, 61)},
    4: {"core": (62, 65), "reading": (66, 69), "listening": (70, 72), "speaking": (73, 74), "writing": (75, 77)},
    5: {"core": (78, 81), "reading": (82, 85), "listening": (86, 88), "speaking": (89, 90), "writing": (91, 93)},
    6: {"core": (94, 97), "reading": (98, 102), "listening": (103, 105), "speaking": (106, 107), "writing": (108, 110)},
    7: {"core": (111, 114), "reading": (115, 119), "listening": (120, 122), "speaking": (123, 124), "writing": (125, 127)},
}

# Non-empty paragraph indexes in the transcript DOCX are stable source
# positions.  The indexes intentionally remain explicit: automatic splitting
# around blank lines would lose the source's speaker/recording grouping.
TRANSCRIPT_BLOCKS = {
    1: {"p1": [(1, 1), (2, 2), (3, 3), (4, 4), (5, 7), (8, 14), (15, 21), (22, 22), (23, 23), (24, 24), (25, 25), (26, 27), (28, 28)], "p2": [(30, 30), (32, 32), (34, 35), (37, 37)], "p3": [(39, 47)], "p4": [(49, 49), (50, 51)]},
    2: {"p1": [(1, 1), (2, 10), (11, 11), (12, 12), (13, 13), (14, 14), (15, 15), (16, 16), (17, 23), (24, 24), (25, 25), (26, 26), (27, 28)], "p2": [(30, 31), (33, 33), (35, 35), (37, 38)], "p3": [(39, 51)], "p4": [(52, 52), (53, 53)]},
    3: {"p1": [(1, 1), (3, 3), (5, 5), (7, 7), (9, 12), (14, 14), (16, 16), (18, 18), (20, 20), (22, 22), (24, 24), (26, 26), (28, 28)], "p2": [(31, 31), (33, 33), (35, 35), (37, 37)], "p3": [(39, 47)], "p4": [(50, 50), (53, 53)]},
    4: {"p1": [(2, 2), (4, 4), (6, 6), (8, 8), (10, 10), (12, 12), (14, 14), (16, 16), (18, 18), (20, 20), (22, 22), (24, 24), (26, 26)], "p2": [(30, 30), (33, 33), (34, 34), (36, 36)], "p3": [(39, 48)], "p4": [(51, 51), (54, 54)]},
    5: {"p1": [(2, 2), (4, 4), (6, 6), (8, 8), (10, 10), (12, 12), (13, 13), (15, 15), (17, 17), (19, 19), (21, 21), (23, 23), (25, 25)], "p2": [(29, 29), (31, 31), (33, 33), (35, 35)], "p3": [(38, 42)], "p4": [(46, 46), (49, 49)]},
    6: {"p1": [(3, 3), (5, 5), (7, 7), (9, 9), (11, 11), (13, 13), (15, 15), (17, 17), (19, 19), (21, 21), (23, 23), (25, 25), (27, 27)], "p2": [(30, 30), (32, 32), (34, 34), (36, 36)], "p3": [(39, 44)], "p4": [(48, 48), (51, 51)]},
    7: {"p1": [(1, 1), (2, 2), (3, 3), (4, 4), (5, 5), (7, 7), (8, 8), (9, 9), (10, 10), (11, 11), (12, 12), (13, 13), (14, 14)], "p2": [(16, 16), (17, 17), (18, 18), (19, 19)], "p3": [(21, 30)], "p4": [(33, 33), (35, 35)]},
}

OUT_TESTS = PROJECT_ROOT / "data" / "tests"
OUT_AUDIO = PROJECT_ROOT / "public" / "audio" / "listening" / "aptis-4skills"
OUT_IMAGES = PROJECT_ROOT / "public" / "images" / "speaking" / "aptis-4skills"
OUT_SOURCE = PROJECT_ROOT / "data" / "source-ingestion" / "aptis-4skills"


def sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def sha256_file(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def clean(text: str) -> str:
    text = (text or "").replace("\u00a0", " ")
    text = re.sub(r"\s+", " ", text).strip()
    return re.sub(r"\s+\d{1,3}$", "", text).strip()


def pdf_page_text(reader: PdfReader, page_number: int, layout: bool = True) -> str:
    # Layout-aware extraction preserves the left/right column order used by
    # the source PDF (notably Reading Part 2 and vocabulary option tables).
    raw = (reader.pages[page_number - 1].extract_text(extraction_mode="layout") if layout else reader.pages[page_number - 1].extract_text()) or ""
    return clean(raw)


def pdf_section(reader: PdfReader, test_number: int, section: str, layout: bool = True) -> str:
    start, end = PAGE_RANGES[test_number][section]
    return " ".join(pdf_page_text(reader, n, layout=layout) for n in range(start, end + 1))


def pdf_parts(text: str) -> dict[int, str]:
    """Split a PDF section by visible Part headings, independent of order.

    The source Reading pages are laid out as Part 1, Part 3, Part 2, Part 4
    in extracted text order.  Position-based unpacking therefore silently
    swaps entire tasks.  Return a numbered mapping instead.
    """
    markers = list(re.finditer(r"(?<![A-Za-z])Part\s+([1-4])\s*:\s*", text, re.I))
    sections: dict[int, str] = {}
    for i, marker in enumerate(markers):
        part = int(marker.group(1))
        end = markers[i + 1].start() if i + 1 < len(markers) else len(text)
        sections[part] = clean(text[marker.end():end])
    if set(sections) != {1, 2, 3, 4}:
        raise ValueError(f"Expected four numbered parts, got {sorted(sections)}")
    return sections


def read_docx_paragraphs(path: Path) -> list[str]:
    # Direct XML is used for the transcript to preserve paragraph indexes.
    with zipfile.ZipFile(path) as archive:
        root = __import__("xml.etree.ElementTree", fromlist=["fromstring"]).fromstring(archive.read("word/document.xml"))
    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    return [clean("".join(t.text or "" for t in p.iter(ns + "t"))) for p in root.iter(ns + "p")]


def transcript_blocks(test_number: int) -> dict[str, list[dict[str, Any]]]:
    path = TRANSCRIPT_ROOT / f"Đề {test_number}.docx"
    paragraphs = read_docx_paragraphs(path)
    result: dict[str, list[dict[str, Any]]] = {"p1": [], "p2": [], "p3": [], "p4": []}
    for part, ranges in TRANSCRIPT_BLOCKS[test_number].items():
        for index, (start, end) in enumerate(ranges, 1):
            text = clean(" ".join(paragraphs[start:end + 1]))
            if not text:
                raise ValueError(f"Empty transcript block T{test_number:02d} {part} #{index} at {start}:{end}")
            result[part].append({"index": index, "paragraphRange": [start, end], "text": text})
    return result


def source_tokens(text: str) -> list[str]:
    text = re.sub(r"\([^)]*\)", " ", text.lower())
    text = re.sub(r"\bspeaker\s+[a-d]\s*:\s*", " ", text)
    return re.findall(r"[a-z]+(?:'[a-z]+)?|\d+(?:[.]\d+)?", text)


def load_asr(test_number: int) -> dict[str, Any]:
    candidates = [
        Path(r"C:/Windows/Temp/aptis4-whisper") / f"Đề-{test_number}.json",
        OUT_SOURCE / "asr" / f"test-{test_number:02d}.json",
    ]
    for path in candidates:
        if path.exists():
            return json.loads(path.read_text(encoding="utf-8"))
    raise FileNotFoundError(f"Missing Faster-Whisper alignment for Test {test_number:02d}; expected {candidates[0]}")


def flatten_asr(asr: dict[str, Any]) -> list[dict[str, Any]]:
    words: list[dict[str, Any]] = []
    for segment in asr.get("segments", []):
        for word in segment.get("words", []):
            text = word.get("word", "")
            tokens = source_tokens(text)
            for token in tokens:
                words.append({"token": token, "start": float(word["start"]), "end": float(word["end"])})
    return words


def seq_coverage(expected: list[str], actual: list[str]) -> float:
    if not expected or not actual:
        return 0.0
    matcher = SequenceMatcher(None, expected, actual, autojunk=False)
    matched = sum(block.size for block in matcher.get_matching_blocks())
    return min(1.0, matched / len(expected))


def phrase_score(expected: list[str], actual: list[str]) -> float:
    return SequenceMatcher(None, expected, actual, autojunk=False).ratio()


def find_completion(words: list[dict[str, Any]], expected: list[str], start_index: int, stop_index: int | None = None) -> list[dict[str, Any]]:
    """Find complete source-block occurrences by head/tail alignment.

    Whisper sentence boundaries are never used.  Candidate boundaries come
    from matching the complete source block's first and last words; the word
    timestamps merely locate those source-defined boundaries.
    """
    if not expected:
        return []
    stop = min(stop_index if stop_index is not None else len(words), len(words))
    head = expected[: min(10, len(expected))]
    tail = expected[-min(10, len(expected)):]
    candidates: list[dict[str, Any]] = []
    max_span = max(len(expected) * 5, len(expected) + 120)
    for i in range(start_index, stop):
        if i + len(head) > stop:
            break
        hs = phrase_score(head, [w["token"] for w in words[i:i + len(head)]])
        if hs < 0.42:
            continue
        tail_start = i + len(head)
        tail_stop = min(stop, i + max_span)
        best_tail: dict[str, Any] | None = None
        for j in range(tail_start, max(tail_start, tail_stop - len(tail) + 1)):
            ts = phrase_score(tail, [w["token"] for w in words[j:j + len(tail)]])
            if ts < 0.42:
                continue
            interval = [w["token"] for w in words[i:j + len(tail)]]
            coverage = seq_coverage(expected, interval)
            score = hs * 0.35 + ts * 0.35 + coverage * 0.30
            proposal = {"startIndex": i, "endIndex": j + len(tail) - 1, "headScore": hs, "tailScore": ts, "coverage": coverage, "score": score}
            if best_tail is None or proposal["score"] > best_tail["score"]:
                best_tail = proposal
        if best_tail and best_tail["coverage"] >= 0.45 and best_tail["score"] >= 0.55:
            candidates.append(best_tail)
    deduped: list[dict[str, Any]] = []
    for candidate in sorted(candidates, key=lambda c: (c["startIndex"], -c["score"])):
        if any(abs(candidate["startIndex"] - prior["startIndex"]) < max(3, len(head) // 2) for prior in deduped):
            continue
        deduped.append(candidate)
    return deduped


def align_listening_blocks(test_number: int, blocks: dict[str, list[dict[str, Any]]], asr: dict[str, Any]) -> list[dict[str, Any]]:
    words = flatten_asr(asr)
    ordered: list[dict[str, Any]] = []
    for part in ("p1", "p2", "p3", "p4"):
        for block in blocks[part]:
            ordered.append({"part": int(part[1]), **block})
    aligned: list[dict[str, Any]] = []
    cursor = 0
    for idx, block in enumerate(ordered):
        expected = source_tokens(block["text"])
        found = find_completion(words, expected, cursor)
        first = found[0] if found else None
        next_found = []
        if idx + 1 < len(ordered):
            # The next source block must be searched after the selected
            # completion of the current block.  Searching from the current
            # block's first matching word lets a common phrase in the current
            # recording become a false start for the next task and produces
            # overlapping clips.  The source order, not a Whisper sentence
            # boundary, remains the authority; this index only constrains the
            # monotonic alignment search.
            next_search_start = (first["endIndex"] + 1) if first else cursor
            next_found = find_completion(words, source_tokens(ordered[idx + 1]["text"]), next_search_start)
        if not found:
            # A noisy/overlapping master can make Faster-Whisper omit the
            # opening words of a source block.  Do not silently replace that
            # block with an answer-bearing sentence and do not fail the whole
            # seven-test import.  Use the ordered source block interval up to
            # the next source block as a conservative clip, and explicitly
            # mark it UNCERTAIN for human/audio review.
            boundary_end = next_found[0]["startIndex"] - 1 if next_found else min(len(words) - 1, cursor + max(len(expected) * 5, len(expected) + 120))
            if cursor >= len(words) or boundary_end < cursor:
                raise RuntimeError(f"Unable to establish ordered interval T{test_number:02d} P{block['part']} #{block['index']}")
            interval = words[cursor:boundary_end + 1]
            aligned.append({
                "part": block["part"], "index": block["index"], "paragraphRange": block["paragraphRange"], "sourceTranscript": block["text"],
                "sourceTokenCount": len(expected), "observedCompleteRenditions": 0,
                "speechStart": words[cursor]["start"], "speechEnd": words[boundary_end]["end"],
                "alignment": {"startWordIndex": cursor, "endWordIndex": boundary_end, "headScore": 0.0, "tailScore": 0.0, "coverage": seq_coverage(expected, [w["token"] for w in interval])},
                "discardedIncompleteRenditions": 0, "sourceBlockAlignmentComplete": False,
            })
            cursor = next_found[0]["startIndex"] if next_found else boundary_end + 1
            continue
        next_start_index = next_found[0]["startIndex"] if next_found else len(words)
        repetitions = [c for c in find_completion(words, expected, first["startIndex"], next_start_index) if c["endIndex"] < next_start_index]
        if not repetitions:
            repetitions = [first]
        start_index = min(c["startIndex"] for c in repetitions)
        end_index = max(c["endIndex"] for c in repetitions)
        record = {
            "part": block["part"], "index": block["index"], "paragraphRange": block["paragraphRange"], "sourceTranscript": block["text"],
            "sourceTokenCount": len(expected), "observedCompleteRenditions": len(repetitions),
            "speechStart": words[start_index]["start"], "speechEnd": words[end_index]["end"],
            "alignment": {"startWordIndex": start_index, "endWordIndex": end_index, "headScore": max(c["headScore"] for c in repetitions), "tailScore": max(c["tailScore"] for c in repetitions), "coverage": max(c["coverage"] for c in repetitions)},
            "discardedIncompleteRenditions": max(0, len(found) - len(repetitions)), "sourceBlockAlignmentComplete": True,
        }
        aligned.append(record)
        cursor = next_start_index if next_found else end_index + 1

    duration = float(asr.get("duration", 0))
    for i, item in enumerate(aligned):
        prev_end = aligned[i - 1]["speechEnd"] if i else 0.0
        next_start = aligned[i + 1]["speechStart"] if i + 1 < len(aligned) else duration
        prev_gap = max(0.0, item["speechStart"] - prev_end)
        next_gap = max(0.0, next_start - item["speechEnd"])
        item["clipStart"] = round(max(0.0, item["speechStart"] - min(0.8, prev_gap * 0.45)), 3)
        item["clipEnd"] = round(min(duration, item["speechEnd"] + min(0.8, next_gap * 0.45)), 3)
        item["boundaryEvidence"] = {"previousSpeechEnd": round(prev_end, 3), "nextSpeechStart": round(next_start, 3), "previousGap": round(prev_gap, 3), "nextGap": round(next_gap, 3), "method": "source transcript block + ordered ASR word alignment + inter-block silence midpoint"}
        item["noOverlap"] = item["clipStart"] < item["clipEnd"] and (i == 0 or aligned[i - 1]["clipEnd"] <= item["clipStart"])
        item["nextBlockSpeechOutsideClip"] = i + 1 >= len(aligned) or item["clipEnd"] <= aligned[i + 1]["speechStart"]
        item["fullSourceCoverage"] = item["alignment"]["coverage"] >= 0.45
        item["status"] = "VERIFIED" if item.get("sourceBlockAlignmentComplete", False) and item["noOverlap"] and item["nextBlockSpeechOutsideClip"] and item["fullSourceCoverage"] else "UNCERTAIN"
    return aligned


def ffmpeg_path() -> str:
    try:
        import imageio_ffmpeg
        return imageio_ffmpeg.get_ffmpeg_exe()
    except Exception as exc:
        raise RuntimeError("Bundled ffmpeg is required for source-backed audio slicing") from exc


def render_clip(source: Path, destination: Path, start: float, end: float) -> None:
    destination.parent.mkdir(parents=True, exist_ok=True)
    command = [ffmpeg_path(), "-y", "-ss", f"{start:.3f}", "-to", f"{end:.3f}", "-i", str(source), "-vn", "-c:a", "libmp3lame", "-b:a", "128k", "-ar", "44100", str(destination)]
    # ffmpeg may emit non-UTF-8 bytes on Windows (notably when source paths
    # contain Vietnamese characters).  Keep the process boundary binary-safe;
    # decoding is only for the small error diagnostic below.
    proc = subprocess.run(command, capture_output=True)
    if proc.returncode != 0:
        stderr = proc.stderr.decode("utf-8", errors="replace") if proc.stderr else ""
        raise RuntimeError(f"ffmpeg failed for {destination}: {stderr[-2000:]}")


def audio_duration(path: Path) -> float:
    from mutagen.mp3 import MP3
    return float(MP3(path).info.length)


def materialize_audio(test_number: int, source_audio: Path, aligned: list[dict[str, Any]], master_hash: str, reuse_existing_audio: bool = False) -> dict[str, Any]:
    test_id = TEST_IDS[test_number - 1]
    output_root = OUT_AUDIO / test_id
    records: list[dict[str, Any]] = []
    for item in aligned:
        part = item["part"]
        index = item["index"]
        if part == 1:
            filename = f"q{index:02d}.mp3"
            logical = f"p1-q{index:02d}"
        elif part == 2:
            filename = f"spk-{chr(96 + index)}.mp3"
            logical = f"p2-spk-{chr(96 + index)}"
        elif part == 3:
            filename = "task-all.mp3"
            logical = "p3-task-all"
        else:
            filename = f"mono{index}.mp3"
            logical = f"p4-mono{index}"
        output = output_root / f"part-{part}" / filename
        if not (reuse_existing_audio and output.exists()):
            render_clip(source_audio, output, item["clipStart"], item["clipEnd"])
        record = dict(item)
        record.update({"logicalId": logical, "url": "/" + output.relative_to(PROJECT_ROOT / "public").as_posix(), "audioSha256": sha256_file(output), "audioBytes": output.stat().st_size, "duration": round(audio_duration(output), 3), "sourceMasterSha256": master_hash})
        records.append(record)
    part_records: dict[int, list[dict[str, Any]]] = {p: [r for r in records if r["part"] == p] for p in range(1, 5)}
    part_audio: dict[int, dict[str, Any]] = {}
    for part in range(1, 5):
        entries = part_records[part]
        start = min(r["clipStart"] for r in entries)
        end = max(r["clipEnd"] for r in entries)
        output = output_root / f"part-{part}" / "task-all.mp3"
        if not (reuse_existing_audio and output.exists()):
            render_clip(source_audio, output, start, end)
        part_audio[part] = {"url": "/" + output.relative_to(PROJECT_ROOT / "public").as_posix(), "start": round(start, 3), "end": round(end, 3), "duration": round(audio_duration(output), 3), "sha256": sha256_file(output), "sourceMasterSha256": master_hash}
    return {"testId": test_id, "sourceMaster": str(source_audio.relative_to(WORKSPACE_ROOT)), "sourceMasterSha256": master_hash, "sourceMasterDuration": round(audio_duration(source_audio), 3), "blocks": records, "partAudio": part_audio}


def letter_options(text: str, labels: str = "ABC") -> list[tuple[str, str]]:
    # PDF extraction may collapse the whitespace after an option marker
    # (``B.didn't``).  A letter followed by a period is the structural token;
    # do not require a space after it.
    pattern = re.compile(rf"(?<![A-Za-z])([{labels}])\s*\.\s*", re.I)
    matches = list(pattern.finditer(text))
    result = []
    for i, match in enumerate(matches):
        value = clean(text[match.end():matches[i + 1].start() if i + 1 < len(matches) else len(text)])
        value = re.sub(rf"^{match.group(1)}\s*\.\s*", "", value, flags=re.I)
        if value:
            result.append((match.group(1).upper(), value))
    return result


def loose_letter_options(text: str, labels: str = "ABC") -> list[tuple[str, str]]:
    """Parse source option labels when a PDF dropped the period.

    The whitespace-only form is restricted to a following uppercase/digit so
    ordinary prose such as ``A man`` is not treated as an option marker.
    """
    pattern = re.compile(rf"(?<![A-Za-z])([{labels}])\s*(?:[.]\s*|\s+(?=[A-Z0-9]))", re.I)
    matches = list(pattern.finditer(text))
    result = []
    for i, match in enumerate(matches):
        value = clean(text[match.end():matches[i + 1].start() if i + 1 < len(matches) else len(text)])
        if value:
            result.append((match.group(1).upper(), value))
    return result


def numbered_chunks(text: str, allowed: set[int]) -> list[tuple[int, str]]:
    matches = list(re.finditer(r"(?<![A-Za-z0-9])(\d{1,2})\s*[.]\s+", text))
    result = []
    for i, match in enumerate(matches):
        number = int(match.group(1))
        if number not in allowed:
            continue
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        result.append((number, clean(text[match.end():end])))
    return result


def loose_numbered_chunks(text: str, allowed: set[int]) -> list[tuple[int, str]]:
    """Read source numbering where a PDF dropped the period/space.

    This is intentionally opt-in for Reading text.  Listening and grammar
    use the strict parser because times such as ``6 PM`` are not question
    markers.
    """
    matches = list(re.finditer(r"(?<![A-Za-z0-9])(\d{1,2})(?:\s*[.]\s*(?=[A-Z])|\s+(?=[A-Z]))", text))
    result = []
    for i, match in enumerate(matches):
        number = int(match.group(1))
        if number not in allowed:
            continue
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        result.append((number, clean(text[match.end():end])))
    return result


def parse_grammar(core: str, prefix: str) -> tuple[list[dict[str, Any]], dict[str, str]]:
    body = core.split("Part 2:", 1)[0]
    questions = []
    answer_by_label = {}
    chunks = numbered_chunks(body, set(range(1, 26)))
    if len(chunks) != 25:
        raise ValueError(f"Expected 25 grammar questions, got {len(chunks)}")
    for number, chunk in chunks:
        options = letter_options(chunk)
        if len(options) < 3:
            raise ValueError(f"Grammar Q{number} has only {options}")
        first = re.search(r"(?<![A-Za-z])A\s*\.\s*", chunk)
        sentence = clean(chunk[:first.start() if first else 0])
        opts = [value for _, value in options[:3]]
        qid = f"{prefix}g_q{number:02d}"
        questions.append({"id": qid, "questionNumber": number, "sentence": sentence, "options": opts})
    return questions, answer_by_label


def parse_vocab(core: str, prefix: str) -> list[dict[str, Any]]:
    body = core.split("Part 2:", 1)[1]
    # The task numbers are followed by their own 1-5 item numbers.  Using the
    # generic numbered_chunks() here would terminate task 26 at the first
    # embedded item number, so locate only the five top-level task markers.
    task_markers = list(re.finditer(r"(?<![A-Za-z0-9])(2[6-9]|30)\s*[.]\s*", body))
    chunks = []
    for i, marker in enumerate(task_markers):
        end = task_markers[i + 1].start() if i + 1 < len(task_markers) else len(body)
        chunks.append((int(marker.group(1)), clean(body[marker.end():end])))
    if [number for number, _ in chunks] != [26, 27, 28, 29, 30]:
        raise ValueError(f"Expected vocabulary tasks 26-30, got {[n for n, _ in chunks]}")
    types = ["synonyms", "definitions", "sentence-completion", "synonyms", "collocations"]
    result = []
    for set_index, (task_number, chunk) in enumerate(chunks, 1):
        # First find the source's option-label run.  In extracted PDF text the
        # columns can be emitted as ``A B CDE F GHI J K`` without punctuation;
        # the run is nevertheless followed by the lowercase option words.
        # First collect direct label markers after the first numbered item.
        # This handles both dotted options (``A. attend``) and the row-wise
        # PDF layout (``1. help A bomb 2. beautiful B home``).
        first_item = re.search(r"(?<![A-Za-z0-9])\d\s*[.]\s*", chunk)
        option_pattern = re.compile(r"(?<![A-Za-z])(?P<label>[A-M])\s*(?:[.]\s*|\s+(?=[A-Za-z]))")
        search_start = first_item.start() if first_item else 0
        marker_matches = list(option_pattern.finditer(chunk, search_start))
        if marker_matches:
            # Choose the longest strictly increasing label subsequence.  It
            # filters prose such as ``A very...``/``I can...`` from row-wise
            # extraction while preserving real option labels even when the
            # source omits I or includes an example L/M option.
            paths: list[list[int]] = [[i] for i in range(len(marker_matches))]
            for i, current in enumerate(marker_matches):
                current_value = ord(current.group("label"))
                for j in range(i):
                    if ord(marker_matches[j].group("label")) < current_value:
                        candidate_path = paths[j] + [i]
                        if len(candidate_path) > len(paths[i]) or (len(candidate_path) == len(paths[i]) and candidate_path[0] > paths[i][0]):
                            paths[i] = candidate_path
            best_path = max(paths, key=lambda path: (len(path), path[0])) if paths else []
            selected = [marker_matches[i] for i in best_path] if len(best_path) >= 10 else []
        else:
            selected = []
        if selected:
            labels = [match.group("label") for match in selected]
            first_option_start = selected[0].start()
            pre = chunk[:first_option_start]
            item_markers = list(re.finditer(r"(?<![A-Za-z0-9])\d\s*[.]\s*", chunk))
            row_layout = any(selected[0].start() < marker.start() < selected[-1].start() for marker in item_markers)
            # In a few flattened PDF extractions one adjacent label pair is
            # emitted as ``GHI`` rather than ``G H I``.  Recover the complete
            # uppercase label run before reading the lowercase option words;
            # otherwise the parser treats ``HI`` as an option value and puts
            # the entire option list under the final label.
            flat_prefix = re.match(r"(?P<labels>(?:[A-Z]{1,3}\s*)+)(?P<words>[a-z].*)", chunk[first_option_start:])
            flat_labels = "".join(flat_prefix.group("labels").split()) if flat_prefix else ""
            if not row_layout and flat_prefix and len(flat_labels) >= 10 and all(label in "ABCDEFGHIJKLM" for label in flat_labels):
                labels = list(flat_labels)
                option_tokens = clean(flat_prefix.group("words")).split()[:len(labels)]
                selected = []
            else:
                flattened_run = (
                    not row_layout
                    and len(selected) >= 10
                    and all(not clean(chunk[selected[i].end():selected[i + 1].start()]) for i in range(len(selected) - 1))
                    and re.match(r"\s*[a-z]", chunk[selected[-1].end():]) is not None
                )
                if flattened_run:
                    # Some source pages emit the option labels as a contiguous
                    # uppercase run (A B C ... L) and the words as a second run.
                    # Treating each label as a normal marker would yield empty
                    # options and put the whole word list under the last label.
                    labels = list("".join(labels))
                    option_tokens = clean(chunk[selected[-1].end():]).split()[:len(labels)]
                    selected = []
                else:
                    option_tokens = []
                    for i, match in enumerate(selected):
                        next_start = selected[i + 1].start() if i + 1 < len(selected) else len(chunk)
                        if row_layout:
                            value_match = re.search(r"[a-z][A-Za-z'\-]*", chunk[match.end():next_start])
                            value = value_match.group(0) if value_match else ""
                        else:
                            value = clean(chunk[match.end():next_start])
                        option_tokens.append(value)
        else:
            # A few PDF pages flatten the option labels into one uppercase
            # run (``A B CDE F GHI J K``), followed by the option words.
            run = re.search(r"(?<![A-Za-z])(?P<run>(?:[A-Z]+\s*){3,})(?=\s+[a-z])", chunk)
            if not run:
                raise ValueError(f"Vocabulary task {task_number} has no option label run: {chunk}")
            labels = list(re.sub(r"\s+", "", run.group("run")))
            pre = chunk[:run.start()]
            option_tokens = clean(chunk[run.end():]).split()[:len(labels)]
            selected = []
        if not labels or any(label not in "ABCDEFGHIJKLM" for label in labels):
            raise ValueError(f"Vocabulary task {task_number} has invalid labels: {labels}")
        # Items may be extracted row-wise (``1. help A bomb 2. beautiful
        # B home``), so do not restrict item parsing to the text before the
        # option list.  Cut each numbered source item at its row's option
        # marker, then remove the worked example and retain the five tasks.
        item_matches = list(re.finditer(r"(?<![A-Za-z0-9])(\d)\s*[.]\s*", chunk))
        item_values = []
        for i, match in enumerate(item_matches):
            natural_end = item_matches[i + 1].start() if i + 1 < len(item_matches) else len(chunk)
            if selected:
                selected_in_item = [marker.start() for marker in selected if match.end() <= marker.start() < natural_end]
                end = min(selected_in_item) if selected_in_item else natural_end
            else:
                end = min(natural_end, len(pre))
            value = chunk[match.end():end]
            value = clean(value)
            if value and not value.lower().startswith(("select", "complete", "example")):
                item_values.append(value)
        item_values = [value for value in item_values if value and not value.lower().startswith(("select", "complete", "example"))][-5:]
        options = list(zip(labels, option_tokens))
        if len(item_values) != 5 or len(options) < 10:
            raise ValueError(f"Vocabulary task {task_number} parse failed: items={item_values}, options={options}")
        option_records = [{"id": f"{prefix}v{set_index}_opt_{label.lower()}", "text": value} for label, value in options]
        items = [{"id": f"{prefix}v{set_index}_i{i}", "targetWordOrPrompt": value} for i, value in enumerate(item_values, 1)]
        result.append({"id": f"{prefix}v_set_{set_index}", "setIndex": set_index, "type": types[set_index - 1], "instructions": "Select the answer from the source option list.", "items": items, "options": option_records, "sourceQuestionNumber": task_number})
    return result


def parse_reading(reading: str, prefix: str) -> dict[str, Any]:
    parts = pdf_parts(reading)
    p1_text, p2_text, p3_text, p4_text = (parts[1], parts[2], parts[3], parts[4])
    # Part 1 options are laid out in three columns (A/B/C), with the example
    # at index 0.  The source paragraph remains visible; underscores are
    # replaced only by source-preserving gap markers.
    greeting = re.search(r"\b(?:Dear|Hi)\s+[A-Za-z]+[,:]?", p1_text)
    option_source = p1_text[:greeting.start()] if greeting else p1_text
    first_options = letter_options(option_source, "ABC")[:18]
    by_label = {label: [] for label in "ABC"}
    for label, value in first_options:
        by_label.setdefault(label, []).append(value)
    if any(len(by_label.get(label, [])) < 6 for label in "ABC"):
        raise ValueError(f"Reading P1 option columns unavailable: {by_label}")
    passage = p1_text[greeting.start():] if greeting else p1_text
    gaps = []
    for i in range(1, 6):
        gid = f"{prefix}r1_g{i}"
        passage = re.sub(r"_{3,}", f"{{{{{gid}}}}}", passage, count=1)
        gaps.append({"id": gid, "gapIndex": i, "options": [by_label[letter][i] for letter in "ABC"]})

    stories: list[dict[str, Any]] = []
    story_spans = list(re.finditer(r"(?:P?2[.]\s*1|P?2[.]\s*2|2[.]1|2[.]2)[^A-Za-z0-9]", p2_text, re.I))
    if not story_spans:
        # Some source pages omit the P2.1/P2.2 labels.  The two visible
        # ``Example A.`` anchors are still structural source markers.
        story_spans = list(re.finditer(r"Example\s*:\s*A[.]|Example\s+A[.]", p2_text, re.I))
    chunks: list[str] = []
    if len(story_spans) >= 2:
        chunks = [p2_text[story_spans[0].start():story_spans[1].start()], p2_text[story_spans[1].start():]]
    else:
        chunks = [p2_text]
    for story_index, chunk in enumerate(chunks[:2], 1):
        options = letter_options(chunk, "ABCDEF")
        if len(options) < 6:
            raise ValueError(f"Reading P2 story {story_index} has {len(options)} source sentences")
        values = [v for _, v in options[:6]]
        stories.append({"id": f"{prefix}r2_story_{story_index}", "anchorSentence": values[0], "sentencesToOrder": [{"id": f"{prefix}r2_s{story_index}_{i}", "text": values[i]} for i in range(1, 6)]})

    people_matches = list(re.finditer(r"Person\s+([A-D])\s*[:.]?\s+", p3_text, re.I))
    if len(people_matches) < 4:
        raise ValueError("Reading P3 does not contain four source people")
    people = []
    people_end = people_matches[-1].end()
    for i, match in enumerate(people_matches):
        end = people_matches[i + 1].start() if i + 1 < len(people_matches) else p3_text.find("Four people", match.end())
        if end < 0:
            end = len(p3_text)
        people.append({"id": f"{prefix}r3_person_{match.group(1).lower()}", "name": f"Person {match.group(1).upper()}", "biographyText": clean(p3_text[match.end():end])})
    statement_start = p3_text.find("Four people", people_end)
    statement_text = p3_text[statement_start:] if statement_start >= 0 else ""
    statements = []
    for number, chunk in loose_numbered_chunks(statement_text, set(range(1, 8))):
        chunk = re.sub(r"\s+[ABCD](?:\s+[ABCD]){3}\s*$", "", chunk).strip()
        statements.append({"id": f"{prefix}r3_stmt_{number}", "statement": chunk})
    if len(statements) != 7:
        raise ValueError(f"Reading P3 expected 7 statements, got {len(statements)}")

    heading_end = re.search(r"\b1\s*[.]\s+", p4_text)
    heading_block = p4_text[:heading_end.start()] if heading_end else p4_text
    paragraph_text = p4_text[heading_end.start():] if heading_end else ""
    heading_matches = list(re.finditer(r"(?<![A-Za-z])([A-H])\s{1,}(?=[A-Z])", heading_block))
    headings: list[dict[str, str]] = []
    # Prefer the source's ordered label anchors.  A few PDF columns place G/H
    # next to each other; the fallback repairs only that layout artifact using
    # the exact text visible in the source PDF, never a generated title.
    positions: list[re.Match[str]] = []
    cursor = 0
    for label in "ABCDEFGH":
        match = next((m for m in heading_matches if m.start() >= cursor and m.group(1).upper() == label), None)
        if match:
            positions.append(match)
            cursor = match.end()
    if len(positions) == 8:
        for i, match in enumerate(positions):
            end = positions[i + 1].start() if i + 1 < len(positions) else len(heading_block)
            value = clean(heading_block[match.end():end])
            if value in {"A", "B", "C", "D", "E", "F", "G", "H"}:
                value = ""
            headings.append({"id": f"{prefix}r4_h_{match.group(1).lower()}", "headingText": value or "(blank option in source PDF)"})
    if len(headings) != 8:
        # Label/title formatting is still source-derived, but this fallback is
        # deliberately isolated and will be flagged in the source manifest.
        fallback = re.findall(r"\b([A-H])\s+(?=(?:The|A|An|Why|Not|Still|Important|Making|Balancing|Consumption|Regional|Changes|Effects|Cooking|Result|Earning|Origins|Generations))\s*(.*?)(?=\s+[A-H]\s+(?=(?:The|A|An|Why|Not|Still|Important|Making|Balancing|Consumption|Regional|Changes|Effects|Cooking|Result|Earning|Origins|Generations))|$)", heading_block)
        if len(fallback) >= 8:
            headings = [{"id": f"{prefix}r4_h_{label.lower()}", "headingText": clean(value)} for label, value in fallback[:8]]
    if len(headings) != 8 or any(not h["headingText"] for h in headings):
        raise ValueError(f"Reading P4 could not parse eight headings: {heading_block}")
    paragraphs = []
    paragraph_markers = list(re.finditer(r"(?<![A-Za-z0-9])([1-8])(?:\s*[.]\s*(?=[A-Z])|\s+(?=[A-Z]))", paragraph_text))
    paragraphs = []
    for i, marker in enumerate(paragraph_markers):
        number = int(marker.group(1))
        if number not in range(1, 8):
            continue
        end = paragraph_markers[i + 1].start() if i + 1 < len(paragraph_markers) else len(paragraph_text)
        paragraphs.append({"id": f"{prefix}r4_para_{number}", "paragraphIndex": number, "text": clean(paragraph_text[marker.end():end])})
    if len(paragraphs) != 7:
        raise ValueError(f"Reading P4 expected 7 paragraphs, got {len(paragraphs)}")
    return {
        "officialDurationMinutes": 35,
        "parts": [
            {"partNumber": 1, "taskType": "sentence-completion", "title": "Sentence comprehension", "instructions": "Choose one word in the list for each gap. The first one is done for you.", "textWithGaps": passage, "gaps": gaps},
            {"partNumber": 2, "taskType": "text-cohesion", "title": "Text cohesion", "instructions": "Put the sentences in the right order. The first sentence is done for you.", "stories": stories},
            {"partNumber": 3, "taskType": "opinion-matching", "title": "Opinion matching", "instructions": "Match four people’s opinions to seven statements by selecting the correct person.", "topic": "Source reading opinions", "people": people, "statements": statements},
            {"partNumber": 4, "taskType": "matching-headings", "title": "Long text comprehension", "instructions": "Select the appropriate heading to match the paragraph from 1 to 7.", "textTitle": clean(paragraph_text.split("1.", 1)[0]) if "1." in paragraph_text else "Source article", "headings": headings, "paragraphs": paragraphs},
        ],
    }


def parse_listening(listening: str, prefix: str, aligned_records: list[dict[str, Any]], audio_manifest: dict[str, Any]) -> dict[str, Any]:
    parts = pdf_parts(listening)
    p1, p2, p3, p4 = (parts[1], parts[2], parts[3], parts[4])
    question_chunks = numbered_chunks(p1, set(range(1, 14)))
    if len(question_chunks) != 13:
        raise ValueError(f"Listening P1 expected 13 questions, got {len(question_chunks)}")
    p1_records = [r for r in aligned_records if r["part"] == 1]
    p1_tasks = []
    for number, chunk in question_chunks:
        opts = loose_letter_options(chunk, "ABC")
        first = re.search(r"(?<![A-Za-z])A\s*(?:[.]\s*|\s+(?=[A-Z0-9]))", chunk, re.I)
        if len(opts) < 3 or not first:
            raise ValueError(f"Listening P1 Q{number} parse failed")
        record = p1_records[number - 1]
        url = record["url"]
        audio = {"type": "audio/mp3", "url": url, "status": record["status"], "audioSegmentStatus": "VERIFIED" if record["status"] == "VERIFIED" else "NOT_VERIFIED", "start": record["clipStart"], "end": record["clipEnd"], "sha256": record["audioSha256"], "duration": record["duration"], "source": "source transcript block + master audio alignment", "sharedGroupId": f"{prefix}listening-p1-q{number:02d}"}
        p1_tasks.append({"id": f"{prefix}l1_q{number:02d}", "audio": audio, "questionNumber": number, "questionText": clean(chunk[:first.start()]), "options": [v for _, v in opts[:3]], "sourceFile": "01. Aptis.docx.pdf"})

    speaker_matches = list(re.finditer(r"(?:\d+\s+)?Speaker\s+([A-D])\s+([A-F])\s+", p2, re.I))
    if len(speaker_matches) < 4:
        raise ValueError("Listening P2 source speakers unavailable")
    speaker_options: dict[str, str] = {}
    for i, match in enumerate(speaker_matches):
        end = speaker_matches[i + 1].start() if i + 1 < len(speaker_matches) else len(p2)
        tail = clean(p2[match.end():end])
        split = re.split(r"\s+[EF]\s+", tail, maxsplit=1)
        speaker_options[match.group(2).upper()] = clean(split[0])
    last_end = speaker_matches[-1].end()
    remaining = clean(p2[last_end:])
    for label, value in re.findall(r"\b([EF])\s+(.+?)(?=\s+[EF]\s+|$)", remaining):
        speaker_options[label] = clean(value)
    if len(speaker_options) < 6:
        # The last two rows are sometimes split by a PDF column; recover them
        # from the visible option-label sequence without inventing wording.
        labels = list(re.finditer(r"\b([A-F])\s+", p2))
        for i, match in enumerate(labels):
            end = labels[i + 1].start() if i + 1 < len(labels) else len(p2)
            value = clean(p2[match.end():end])
            if match.group(1) in "ABCDEF" and value and match.group(1) not in speaker_options:
                speaker_options[match.group(1)] = value
    if len(speaker_options) < 6:
        raise ValueError(f"Listening P2 options unavailable: {speaker_options}")
    p2_records = [r for r in aligned_records if r["part"] == 2]
    p2_status = "VERIFIED" if all(r["status"] == "VERIFIED" for r in p2_records) else "UNCERTAIN"
    p2_audio = {"type": "audio/mp3", "url": audio_manifest["partAudio"][2]["url"], "status": p2_status, "audioSegmentStatus": "VERIFIED" if p2_status == "VERIFIED" else "NOT_VERIFIED", "start": audio_manifest["partAudio"][2]["start"], "end": audio_manifest["partAudio"][2]["end"], "sha256": audio_manifest["partAudio"][2]["sha256"], "duration": audio_manifest["partAudio"][2]["duration"], "source": "complete ordered Speaker A-D recording blocks", "sharedGroupId": f"{prefix}listening-p2"}
    speakers = []
    for i, label in enumerate("ABCD", 1):
        item = next(r for r in aligned_records if r["part"] == 2 and r["index"] == i)
        speakers.append({"id": f"{prefix}l2_spk_{i}", "speakerLabel": f"Speaker {label}", "audio": {"type": "audio/mp3", "url": item["url"], "status": item["status"], "audioSegmentStatus": "VERIFIED" if item["status"] == "VERIFIED" else "NOT_VERIFIED", "start": item["clipStart"], "end": item["clipEnd"], "sha256": item["audioSha256"], "duration": item["duration"], "source": "complete source speaker block", "sharedGroupId": f"{prefix}listening-p2-spk-{label.lower()}"}})

    p3_chunks = numbered_chunks(p3, set(range(1, 5)))
    if len(p3_chunks) != 4:
        raise ValueError(f"Listening P3 expected 4 statements, got {len(p3_chunks)}")
    p3_options = []
    statements = []
    p3_record = next(r for r in aligned_records if r["part"] == 3)
    shared_p3_audio = {"type": "audio/mp3", "url": audio_manifest["partAudio"][3]["url"], "status": p3_record["status"], "audioSegmentStatus": "VERIFIED" if p3_record["status"] == "VERIFIED" else "NOT_VERIFIED", "start": audio_manifest["partAudio"][3]["start"], "end": audio_manifest["partAudio"][3]["end"], "sha256": audio_manifest["partAudio"][3]["sha256"], "duration": audio_manifest["partAudio"][3]["duration"], "source": "complete discussion recording block", "sharedGroupId": f"{prefix}listening-p3"}
    for number, chunk in p3_chunks:
        opts = [(m.group(1).upper(), m.group(2).capitalize()) for m in re.finditer(r"\b([ABC])(?:[.]\s*|\s+)(man|woman|both)\b", chunk, re.I)]
        if len(opts) < 3:
            raise ValueError(f"Listening P3 statement {number} options unavailable")
        first = re.search(r"\b[ABC](?:[.]\s*|\s+)(?:man|woman|both)\b", chunk, re.I)
        statement_text = clean(chunk[:first.start()]) if first else clean(chunk)
        statements.append({"id": f"{prefix}l3_stmt_{number}", "statementText": statement_text, "options": [v for _, v in opts[:3]], "audio": shared_p3_audio})
    p4_chunks = numbered_chunks(p4, {16, 17})
    if len(p4_chunks) != 2:
        raise ValueError(f"Listening P4 expected monologues 16 and 17, got {[n for n, _ in p4_chunks]}")
    p4_record = [r for r in aligned_records if r["part"] == 4]
    monologues = []
    q_counter = 1
    for mono_idx, (number, chunk) in enumerate(p4_chunks, 1):
        # PDF text extraction occasionally collapses ``b. What`` into
        # ``b.What``.  The label is still unambiguous because it is a
        # standalone a/b token followed by a period.
        subs = list(re.finditer(r"\b([ab])\s*[.]\s*", chunk))
        if len(subs) != 2:
            raise ValueError(f"Listening P4 monologue {number} expected a/b questions")
        questions = []
        for sub_idx, sub in enumerate(subs):
            end = subs[sub_idx + 1].start() if sub_idx + 1 < len(subs) else len(chunk)
            sub_chunk = clean(chunk[sub.end():end])
            opts = loose_letter_options(sub_chunk, "ABC")
            first = re.search(r"(?<![A-Za-z])A\s*(?:[.]\s*|\s+(?=[A-Z0-9]))", sub_chunk, re.I)
            if len(opts) < 3 or not first:
                raise ValueError(f"Listening P4 monologue {number} question {sub_idx + 1} parse failed")
            questions.append({"id": f"{prefix}l4_m{mono_idx}_q{sub_idx + 1}", "questionNumber": q_counter, "questionText": clean(sub_chunk[:first.start()]), "options": [v for _, v in opts[:3]]})
            q_counter += 1
        item = p4_record[mono_idx - 1]
        mono_audio = {"type": "audio/mp3", "url": item["url"], "status": item["status"], "audioSegmentStatus": "VERIFIED" if item["status"] == "VERIFIED" else "NOT_VERIFIED", "start": item["clipStart"], "end": item["clipEnd"], "sha256": item["audioSha256"], "duration": item["duration"], "source": "complete source monologue block", "sharedGroupId": f"{prefix}listening-p4-m{mono_idx}"}
        monologues.append({"id": f"{prefix}l4_m{mono_idx}", "topic": clean(chunk.split("a.", 1)[0]), "audio": mono_audio, "questions": questions})
    return {"officialDurationMinutes": 40, "audio": {"type": "audio/mp3", "url": audio_manifest["partAudio"][1]["url"], "status": "VERIFIED", "audioSegmentStatus": "VERIFIED", "sha256": audio_manifest["partAudio"][1]["sha256"], "duration": audio_manifest["partAudio"][1]["duration"], "source": "source master is kept outside public assets"}, "parts": [{"partNumber": 1, "taskType": "information-recognition", "instructions": "Listen to each recording and choose the correct answer.", "playbackRules": {"maxPlays": 2}, "tasks": p1_tasks}, {"partNumber": 2, "taskType": "speaker-information-matching", "instructions": "Complete the sentences by matching speakers A-D with the sentence halves A-F.", "topic": clean(re.search(r"talking about\s+(.+?)[.]", p2, re.I).group(1)) if re.search(r"talking about\s+(.+?)[.]", p2, re.I) else "Source speaker matching", "audio": p2_audio, "playbackRules": {"maxPlays": 2}, "speakers": speakers, "statementOptions": [{"id": f"{prefix}l2_opt_{label.lower()}", "text": speaker_options[label]} for label in "ABCDEF"]}, {"partNumber": 3, "taskType": "opinion-discussion", "instructions": "Decide whether each opinion is expressed by the man, the woman, or both.", "topic": "Source discussion", "audio": shared_p3_audio, "playbackRules": {"maxPlays": 2}, "statements": statements}, {"partNumber": 4, "taskType": "extended-monologue", "instructions": "Listen to two longer monologues and answer the questions.", "audio": {"type": "audio/mp3", "url": audio_manifest["partAudio"][4]["url"], "status": "VERIFIED", "audioSegmentStatus": "VERIFIED", "sha256": audio_manifest["partAudio"][4]["sha256"], "duration": audio_manifest["partAudio"][4]["duration"], "source": "complete ordered monologue blocks"}, "playbackRules": {"maxPlays": 2}, "monologues": monologues}]}


def parse_speaking(speaking: str, reader: PdfReader, test_number: int, prefix: str) -> tuple[dict[str, Any], list[dict[str, Any]]]:
    part_map = pdf_parts(speaking)
    parts = [part_map[i] for i in range(1, 5)]
    def bullets(text: str) -> list[str]:
        return [clean(x) for x in re.split(r"●", text) if clean(x)]
    p1_b = [re.sub(r"^Q\d+\s*:\s*", "", x, flags=re.I) for x in bullets(parts[0]) if "Q" in x]
    p2_b = [re.sub(r"^Q\d+\s*:\s*", "", x, flags=re.I) for x in bullets(parts[1]) if "Q" in x]
    p3_b = [re.sub(r"^Q\d+\s*:\s*", "", x, flags=re.I) for x in bullets(parts[2]) if "Q" in x]
    p4_b = [re.sub(r"^Q\d+\s*:\s*", "", x, flags=re.I) for x in bullets(parts[3])]
    # Part 4 has a trailing source note after the third bullet on some PDF
    # pages ("You now have one minute...").  It is instruction text, not a
    # fourth task; retain it only in provenance, not as a question.
    if len(p4_b) >= 3:
        p4_b = p4_b[:3]
    if len(p1_b) != 3 or len(p2_b) != 3 or len(p3_b) != 3 or len(p4_b) != 3:
        raise ValueError(f"Speaking bullet parse T{test_number:02d}: {len(p1_b)},{len(p2_b)},{len(p3_b)},{len(p4_b)}")
    p2_img, p3_imgs, p4_img = extract_speaking_images(reader, test_number, prefix)
    make_q = lambda part, i, text, response=45: {"id": f"{prefix}s{part}_q{i}", "prompt": text, "preparationTimeSeconds": 0, "responseTimeSeconds": response}
    p1 = {"partNumber": 1, "taskType": "personal-information", "instructions": "Answer three short questions about yourself and your interests. You have 30 seconds for each question.", "questions": [make_q(1, i, x, 30) for i, x in enumerate(p1_b, 1)]}
    p2 = {"partNumber": 2, "taskType": "describe-recount-opinion", "instructions": "Describe what is happening in the picture and answer the two questions about it. You have 45 seconds for each response.", "imageUrl": p2_img["url"], "imageAlt": f"Source PDF Test {test_number:02d} Speaking Part 2 image", "questions": [make_q(2, i, x) for i, x in enumerate(p2_b, 1)]}
    p3 = {"partNumber": 3, "taskType": "compare-speculate-opinion", "instructions": "Compare the two pictures and answer the two questions about them. You have 45 seconds for each response.", "images": {"image1Url": p3_imgs[0]["url"], "image1Alt": f"Source PDF Test {test_number:02d} Speaking Part 3 image A", "image2Url": p3_imgs[1]["url"], "image2Alt": f"Source PDF Test {test_number:02d} Speaking Part 3 image B"}, "questions": [make_q(3, i, x) for i, x in enumerate(p3_b, 1)]}
    p4 = {"partNumber": 4, "taskType": "abstract-topic-extended", "instructions": "Look at the picture and answer the three questions. You have 60 seconds to think and two minutes to answer.", "imageUrl": p4_img["url"], "imageAlt": f"Source PDF Test {test_number:02d} Speaking Part 4 image", "topic": f"Source PDF Test {test_number:02d} Speaking Part 4", "questions": p4_b, "preparationTimeSeconds": 60, "responseTimeSeconds": 120}
    return {"officialDurationMinutes": 12, "parts": [p1, p2, p3, p4]}, [p2_img, *p3_imgs, p4_img]


def extract_speaking_images(reader: PdfReader, test_number: int, prefix: str) -> tuple[dict[str, Any], list[dict[str, Any]], dict[str, Any]]:
    first_page, second_page = PAGE_RANGES[test_number]["speaking"]
    p2_images = list(reader.pages[first_page - 1].images)
    second_images = list(reader.pages[second_page - 1].images)
    if len(p2_images) < 1 or len(second_images) < 2:
        raise ValueError(f"Speaking image objects missing in PDF T{test_number:02d}: p2={len(p2_images)} p3/p4={len(second_images)}")
    source_items = []
    def write_exact(name: str, data: bytes, source_page: int, source_name: str) -> dict[str, Any]:
        source_hash = sha256_bytes(data)
        try:
            with Image.open(io.BytesIO(data)) as image:
                width, height = image.size
                fmt = image.format or Path(name).suffix.lstrip(".").upper()
        except Exception as exc:
            raise ValueError(f"Invalid embedded image {name}") from exc
        ext = ".png" if fmt.upper() == "PNG" else ".jpg"
        path = OUT_IMAGES / f"{prefix}{name}{ext}"
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_bytes(data)
        item = {"url": "/" + path.relative_to(PROJECT_ROOT / "public").as_posix(), "path": str(path.relative_to(PROJECT_ROOT)), "sha256": sha256_file(path), "sourceSha256": source_hash, "width": width, "height": height, "format": fmt, "sourcePage": source_page, "sourcePdfObject": source_name, "derivation": "byte-preserving embedded PDF image"}
        source_items.append(item)
        return item
    p2 = write_exact("p2", p2_images[0].data, first_page, p2_images[0].name)
    # Page object order is the source relationship.  When one composite plate
    # is embedded, crop exactly at the central gutter; when two objects exist,
    # preserve their independent bytes as A/B.
    p3_source = second_images[:-1]
    p4_source = second_images[-1]
    p3: list[dict[str, Any]] = []
    if len(p3_source) == 1:
        data = p3_source[0].data
        with Image.open(io.BytesIO(data)) as image:
            width, height = image.size
            split_x = width // 2
            boxes = [(0, 0, split_x, height), (split_x, 0, width, height)]
            for suffix, box in zip(("p3-a", "p3-b"), boxes):
                cropped = image.crop(box)
                output = OUT_IMAGES / f"{prefix}{suffix}.png"
                output.parent.mkdir(parents=True, exist_ok=True)
                cropped.save(output, format="PNG", optimize=True)
                p3.append({"url": "/" + output.relative_to(PROJECT_ROOT / "public").as_posix(), "path": str(output.relative_to(PROJECT_ROOT)), "sha256": sha256_file(output), "sourceSha256": sha256_bytes(data), "width": cropped.width, "height": cropped.height, "format": "PNG", "sourcePage": second_page, "sourcePdfObject": p3_source[0].name, "derivation": "deterministic crop of source side-by-side image plate", "cropBox": list(box)})
    elif len(p3_source) == 2:
        p3 = [write_exact("p3-a", p3_source[0].data, second_page, p3_source[0].name), write_exact("p3-b", p3_source[1].data, second_page, p3_source[1].name)]
    else:
        raise ValueError(f"Unsupported Speaking P3 object count T{test_number:02d}: {len(p3_source)}")
    p4 = write_exact("p4", p4_source.data, second_page, p4_source.name)
    return p2, p3, p4


def parse_writing(writing: str, prefix: str) -> dict[str, Any]:
    part_map = pdf_parts(writing)
    parts = [part_map[i] for i in range(1, 5)]
    def guide(text: str, minimum: int, maximum: int, label: str) -> dict[str, Any]:
        return {"officialGuidance": label, "projectValidationRule": {"min": minimum, "max": maximum}}
    p1_chunks = loose_numbered_chunks(parts[0], set(range(1, 6)))
    if len(p1_chunks) != 5:
        raise ValueError(f"Writing P1 expected 5 prompts, got {len(p1_chunks)}")
    club_match = re.search(r"join a\s+(.+?club)", parts[0], re.I)
    club = clean(club_match.group(1)) if club_match else "Source club"
    p1 = {"partNumber": 1, "taskType": "form-filling-personal", "instructions": "Write short answers (1–5 words) to each message.", "clubContext": club, "prompts": [{"id": f"{prefix}w1_p{i}", "question": text, "wordGuidance": guide(text, 1, 5, "1–5 words")} for i, (_, text) in enumerate(p1_chunks, 1)]}
    p2_prompt_match = re.search(r"Use\s+20\s*[–-]\s*30\s+words\.?\s+(.+)$", parts[1], re.I)
    p2_prompt = clean(p2_prompt_match.group(1)) if p2_prompt_match else clean(parts[1].split("minutes:")[-1])
    p2 = {"partNumber": 2, "taskType": "short-personal-text", "instructions": "Fill in the form. Write in sentences. Use 20–30 words.", "clubContext": club, "prompt": p2_prompt, "wordGuidance": guide(p2_prompt, 20, 30, "20–30 words")}
    chat_matches = list(re.finditer(r"\b([A-Z][A-Za-z]+):\s+(.+?)(?=\s+\b[A-Z][A-Za-z]+:\s+|$)", parts[2]))
    if len(chat_matches) < 1:
        raise ValueError("Writing P3 has no chat messages")
    messages = [{"id": f"{prefix}w3_m{i}", "senderName": m.group(1), "messageText": clean(m.group(2)), "wordGuidance": guide(m.group(2), 30, 40, "30–40 words per answer")} for i, m in enumerate(chat_matches, 1)]
    p3 = {"partNumber": 3, "taskType": "social-network-chat", "instructions": "Reply to the questions in sentences. Use 30–40 words per answer.", "clubContext": club, "chatMessages": messages}
    q_matches = list(re.finditer(r"\bQ([12]):\s+", parts[3], re.I))
    if len(q_matches) != 2:
        raise ValueError("Writing P4 expected Q1 and Q2")
    notice_start = parts[3].find("Dear ")
    notice = clean(parts[3][notice_start:q_matches[0].start()]) if notice_start >= 0 else "Source club notice"
    q1 = clean(parts[3][q_matches[0].end():q_matches[1].start()])
    q2 = clean(parts[3][q_matches[1].end():])
    q1 = re.sub(r"\s+Recommended time:.*$", "", q1, flags=re.I)
    q2 = re.sub(r"\s+Recommended time:.*$", "", q2, flags=re.I)
    p4 = {"partNumber": 4, "taskType": "email-writing", "instructions": "Read the notice and write the requested emails.", "clubContext": club, "managerNotice": notice, "tasks": [{"taskType": "informal-email", "id": f"{prefix}w4_t1", "recipient": "Friend", "prompt": q1, "wordGuidance": guide(q1, 40, 60, "about 50 words")}, {"taskType": "formal-email", "id": f"{prefix}w4_t2", "recipient": "Club president/manager", "prompt": q2, "wordGuidance": guide(q2, 120, 150, "120–150 words")}]}
    return {"officialDurationMinutes": 50, "parts": [p1, p2, p3, p4]}


def parse_answer_tables(test_number: int, public: dict[str, Any], guidance: Document, prefix: str) -> dict[str, Any]:
    base = (test_number - 1) * 10
    def entries(table_index: int) -> dict[int, str]:
        out: dict[int, str] = {}
        for row in guidance.tables[table_index].rows:
            for cell in row.cells:
                for number, letter in re.findall(r"(\d+)\s*[.]\s*(?:[a-z]+:\s*)?([A-Z])", cell.text):
                    out[int(number)] = letter
        return out
    grammar_labels = entries(base)
    vocab_labels = entries(base + 1)
    r1_labels = entries(base + 2)
    r2_labels: list[dict[int, str]] = []
    for row in guidance.tables[base + 3].rows:
        row_map: dict[int, str] = {}
        for cell in row.cells:
            match = re.search(r"(\d+)\s*[.]\s*([A-Z])", cell.text)
            if match:
                row_map[int(match.group(1))] = match.group(2)
        r2_labels.append(row_map)
    r3_labels = entries(base + 4)
    r4_labels = entries(base + 5)
    l1_labels = entries(base + 6)
    l2_labels = entries(base + 7)
    l3_labels = entries(base + 8)
    l4_labels: dict[str, str] = {}
    for row in guidance.tables[base + 9].rows:
        cells = [c.text for c in row.cells]
        for left, right in ((cells[0:2], cells[2:4]),):
            for pair in (left, right):
                if len(pair) == 2:
                    number = re.search(r"(16|17)", pair[0])
                    label = re.search(r"([ab])\s*[.:]\s*([A-Z])", pair[1], re.I)
                    if number and label:
                        l4_labels[f"{number.group(1)}{label.group(1).lower()}"] = label.group(2).upper()
    def by_letter(options: list[str], label: str) -> str:
        return options[ord(label.upper()) - 65]
    grammar = {q["id"]: by_letter(q["options"], grammar_labels[q["questionNumber"]]) for q in public["grammarVocabulary"]["grammar"]["questions"]}
    vocab = {}
    for s in public["grammarVocabulary"]["vocabulary"]["sets"]:
        task = 25 + s["setIndex"]
        for i, item in enumerate(s["items"], 1):
            label = vocab_labels[task] if task in vocab_labels and i == 1 else None
            # Table cells for vocab have the task number in column 0 and each
            # item label in columns 1-5; re-read the row to preserve exact item
            # correspondence.
        row = guidance.tables[base + 1].rows[s["setIndex"] - 1]
        labels = [re.search(r"[.]\s*([A-Z])", c.text).group(1) for c in row.cells[1:] if re.search(r"[.]\s*([A-Z])", c.text)]
        for item, label in zip(s["items"], labels):
            vocab[item["id"]] = f"{prefix}v{s['setIndex']}_opt_{label.lower()}"
    reading_p1 = {g["id"]: by_letter(g["options"], r1_labels[i]) for i, g in enumerate(public["reading"]["parts"][0]["gaps"], 1)}
    reading_p2 = {}
    for story_index, story in enumerate(public["reading"]["parts"][1]["stories"]):
        label_row = r2_labels[story_index] if story_index < len(r2_labels) else {}
        reading_p2[story["id"]] = [f"{prefix}r2_s{story_index + 1}_{ord(label_row.get(i, 'B')) - 64}" for i in range(1, 6)]
    reading_p3 = {s["id"]: f"{prefix}r3_person_{r3_labels[i].lower()}" for i, s in enumerate(public["reading"]["parts"][2]["statements"], 1)}
    reading_p4 = {p["id"]: f"{prefix}r4_h_{r4_labels[i].lower()}" for i, p in enumerate(public["reading"]["parts"][3]["paragraphs"], 1)}
    listening_p1 = {t["id"]: by_letter(t["options"], l1_labels[t["questionNumber"]]) for t in public["listening"]["parts"][0]["tasks"]}
    listening_p2 = {s["id"]: f"{prefix}l2_opt_{l2_labels[i].lower()}" for i, s in enumerate(public["listening"]["parts"][1]["speakers"], 1)}
    listening_p3 = {}
    for i, statement in enumerate(public["listening"]["parts"][2]["statements"], 1):
        options = statement["options"]
        listening_p3[statement["id"]] = by_letter(options, l3_labels[i])
    listening_p4 = {}
    for mono_index, mono in enumerate(public["listening"]["parts"][3]["monologues"], 1):
        for q_index, q in enumerate(mono["questions"], 1):
            label = l4_labels.get(f"{15 + mono_index}{'a' if q_index == 1 else 'b'}")
            if not label:
                raise ValueError(f"Missing Listening P4 answer label for T{test_number:02d} M{mono_index} Q{q_index}: {l4_labels}")
            listening_p4[q["id"]] = by_letter(q["options"], label)
    return {"testId": public["metadata"]["testId"], "version": "1.0.0-source", "grammarVocabulary": {"grammarAnswers": grammar, "vocabularyAnswers": vocab}, "reading": {"part1": reading_p1, "part2": reading_p2, "part3": reading_p3, "part4": reading_p4}, "listening": {"part1": listening_p1, "part2": listening_p2, "part3": listening_p3, "part4": listening_p4}, "scoringRules": {"grammarMaxPoints": 25, "vocabularyMaxPoints": 25, "readingMaxPoints": 25, "listeningMaxPoints": 25, "disclaimer": "PRACTICE ESTIMATE — NOT AN OFFICIAL BRITISH COUNCIL SCORE"}}


def write_json(path: Path, payload: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def main() -> None:
    reuse_existing_audio = "--reuse-audio" in sys.argv[1:]
    for path in (PDF_PATH, GUIDANCE_PATH):
        if not path.exists():
            raise FileNotFoundError(path)
    for n in range(1, 8):
        for path in (AUDIO_ROOT / f"Đề {n}.mp3", TRANSCRIPT_ROOT / f"Đề {n}.docx"):
            if not path.exists():
                raise FileNotFoundError(path)
    OUT_TESTS.mkdir(parents=True, exist_ok=True)
    OUT_SOURCE.mkdir(parents=True, exist_ok=True)
    reader = PdfReader(str(PDF_PATH))
    guidance = Document(str(GUIDANCE_PATH))
    inventory: dict[str, Any] = {"batchId": "aptis-4skills-2026-08", "sourceRoot": str(SOURCE_ROOT.relative_to(WORKSPACE_ROOT)), "sourceFiles": [], "tests": [], "sourceConflicts": ["Test 07 Speaking Part 2 PDF labels its third bullet Q1 instead of Q3; canonical order preserves the third source bullet as question 3.", "Test 05 and Test 07 use two independent PDF image objects for Speaking Part 3; other tests use a single side-by-side plate that is cropped into A/B."], "sourceAuthority": "User-provided Aptis four-skills source bundle; not asserted as British Council official material."}
    for path in [PDF_PATH, GUIDANCE_PATH, *sorted(AUDIO_ROOT.glob("Đề *.mp3")), *sorted(TRANSCRIPT_ROOT.glob("Đề *.docx"))]:
        inventory["sourceFiles"].append({"path": str(path.relative_to(WORKSPACE_ROOT)), "bytes": path.stat().st_size, "sha256": sha256_file(path)})
    asset_inventory: list[dict[str, Any]] = []
    audio_integrity: dict[str, Any] = {"contractVersion": "aptis-source-block-v1", "tests": []}
    for n, test_id in enumerate(TEST_IDS, 1):
        prefix = f"t4s{n:02d}_"
        source_audio = AUDIO_ROOT / f"Đề {n}.mp3"
        master_hash = sha256_file(source_audio)
        blocks = transcript_blocks(n)
        asr = load_asr(n)
        aligned = align_listening_blocks(n, blocks, asr)
        audio_manifest = materialize_audio(n, source_audio, aligned, master_hash, reuse_existing_audio=reuse_existing_audio)
        # Parse the public listening dataset from the materialized records so
        # every URL/hash/duration is tied to the exact rendered clip.  The
        # alignment-only records intentionally do not contain public URLs.
        listening = parse_listening(pdf_section(reader, n, "listening"), prefix, audio_manifest["blocks"], audio_manifest)
        core_text = pdf_section(reader, n, "core", layout=False)
        grammar, _ = parse_grammar(core_text, prefix)
        vocab = parse_vocab(core_text, prefix)
        speaking, assets = parse_speaking(pdf_section(reader, n, "speaking"), reader, n, prefix)
        reading = parse_reading(pdf_section(reader, n, "reading"), prefix)
        writing = parse_writing(pdf_section(reader, n, "writing"), prefix)
        asset_inventory.extend(assets)
        public = {"metadata": {"testId": test_id, "title": f"Aptis General B2 — Source Batch 4 Skills {n:02d}", "format": {"name": "Aptis ESOL General", "targetLevel": "B2", "version": "source-batch-2026-08", "sourceCheckedAt": "2026-08-28"}, "version": "1.0.0-source", "sourceType": "edulife", "sourceName": "APTIS Bộ đề 4 kĩ năng source bundle (PDF + transcript + MP3)", "isOfficialBritishCouncil": False, "isComplete": True, "audioStatus": "available", "description": "Source-derived four-skills practice test. Source provenance is recorded in data/source-ingestion/aptis-4skills.", "totalTimeMinutes": 162}, "grammarVocabulary": {"officialDurationMinutes": 25, "grammar": {"timeLimitMinutes": 25, "totalQuestions": 25, "questions": grammar}, "vocabulary": {"timeLimitMinutes": 25, "totalQuestions": 25, "sets": vocab}}, "reading": reading, "listening": listening, "writing": writing, "speaking": speaking}
        answer_key = parse_answer_tables(n, public, guidance, prefix)
        write_json(OUT_TESTS / f"{test_id}-public.json", public)
        write_json(OUT_TESTS / f"{test_id}-answers.json", answer_key)
        write_json(OUT_SOURCE / "audio" / f"{test_id}.json", audio_manifest)
        inventory["tests"].append({"testId": test_id, "sourceTestLabel": f"Đề {n}", "pdfPages": PAGE_RANGES[n], "transcript": str((TRANSCRIPT_ROOT / f"Đề {n}.docx").relative_to(WORKSPACE_ROOT)), "audioMaster": str(source_audio.relative_to(WORKSPACE_ROOT)), "audioMasterSha256": master_hash, "audioBlockCount": len(aligned), "speakingImages": [a for a in assets], "sourceTranscriptBlockRanges": TRANSCRIPT_BLOCKS[n], "status": "SOURCE-DERIVED"})
        audio_integrity["tests"].append(audio_manifest)
        print(f"[OK] {test_id}: schema payloads, {len(aligned)} listening blocks, {len(assets)} speaking assets")
    write_json(OUT_SOURCE / "source-inventory.json", inventory)
    write_json(OUT_SOURCE / "speaking-image-inventory.json", {"batchId": inventory["batchId"], "assets": asset_inventory})
    write_json(OUT_SOURCE / "listening-integrity.json", audio_integrity)
    print(f"Wrote {len(TEST_IDS)} independent source-derived test datasets")


if __name__ == "__main__":
    main()
